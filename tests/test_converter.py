from pathlib import Path

import pytest

from scripts import convert_inputs


def test_convert_txt_reads_utf8(tmp_path: Path) -> None:
    sample = tmp_path / "sample.txt"
    sample.write_text("Linha 1\n\nLinha 2", encoding="utf-8")

    text = convert_inputs.convert_txt(sample)

    assert "Linha 1" in text
    assert "Linha 2" in text
    assert "\n\n" in text  # parágrafos separados


def test_convert_docx_to_paragraphs(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")  # noqa: F841 - exercised by conversion
    from docx import Document

    sample = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Primeiro parágrafo")
    document.add_paragraph("Segundo parágrafo")
    document.save(sample)

    text = convert_inputs.convert_docx(sample)

    assert "Primeiro parágrafo" in text
    assert "Segundo parágrafo" in text
    assert "\n\n" in text


def test_convert_file_writes_output(tmp_path: Path) -> None:
    sample = tmp_path / "sample.txt"
    sample.write_text("Alpha\n\nBeta", encoding="utf-8")
    output = tmp_path / "out.md"

    text = convert_inputs.convert_file(sample, output)

    assert output.read_text(encoding="utf-8") == text
    assert "Alpha" in text
